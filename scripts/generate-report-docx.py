#!/usr/bin/env python3
"""Genera il report privacy AMCO Service in formato Word (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# Percorso output
output_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
output_path = os.path.join(output_dir, "docs", "Report Verifica Privacy Sito Web - AMCO Service - Marzo 2026.docx")

doc = Document()

# ============================================
# STILI
# ============================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.15

# Heading styles
for level in [1, 2, 3]:
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
    heading_style.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    heading_style.paragraph_format.space_after = Pt(8)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


def add_table(doc, headers, rows, col_widths=None):
    """Aggiunge una tabella formattata"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Sfondo header
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '384b5a',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

    # Rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.size = Pt(10)
            # Alternating row colors
            if row_idx % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'f1f5f9',
                    qn('w:val'): 'clear',
                })
                shading.append(shading_elm)

    # Larghezze colonne
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # Spaziatura
    return table


def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ============================================
# INTESTAZIONE
# ============================================
# Titolo
title = doc.add_heading('Verifica Compliance Privacy Sito Web', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AMCO Service SRL — www.amcovr.it')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x5a, 0x98, 0x1d)
run.bold = True

doc.add_paragraph()

# Info box
info_table = doc.add_table(rows=3, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('Data', '9 marzo 2026'),
    ('A cura di', 'Midas Touch SRL — Servizio Amministratore Protetto'),
    ('Destinatario', 'AMCO Service SRL, Via Francesco Morosini 7, 37138 Verona (VR)'),
]
for i, (label, value) in enumerate(info_data):
    cell_label = info_table.rows[i].cells[0]
    cell_label.text = label
    cell_label.paragraphs[0].runs[0].bold = True
    cell_label.paragraphs[0].runs[0].font.size = Pt(10)
    cell_label.width = Cm(3.5)
    shading = cell_label._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'f1f5f9', qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

    cell_value = info_table.rows[i].cells[1]
    cell_value.text = value
    cell_value.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ============================================
# PREMESSA
# ============================================
doc.add_heading('Premessa', level=2)
doc.add_paragraph(
    'Il presente documento riporta le osservazioni emerse da una verifica del sito web '
    'www.amcovr.it sotto il profilo della normativa privacy vigente (Regolamento UE 2016/679 — GDPR).'
)
doc.add_paragraph(
    'Le raccomandazioni sono fornite a titolo collaborativo nell\'ambito del rapporto di consulenza '
    'in essere. Il sito web aziendale resta di esclusiva titolarità e responsabilità di AMCO Service SRL; '
    'l\'implementazione delle eventuali modifiche è a carico dello studio e del proprio fornitore tecnico.'
)

# ============================================
# 1. QUADRO GENERALE
# ============================================
doc.add_heading('1. Quadro generale', level=2)
doc.add_paragraph(
    'Il sito www.amcovr.it è un sito vetrina realizzato in WordPress. Contiene pagine informative '
    'sull\'attività dello studio, form di contatto e richiesta preventivo con protezione Google reCAPTCHA, '
    'accesso all\'area condomini (tramite piattaforma esterna miocondominio.eu), e due pagine dedicate alla privacy.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Elementi positivi già presenti:')

positives = [
    'Informativa privacy e cookie policy linkate nel footer',
    'Dati aziendali visibili (P.IVA, REA, riferimento L. 4/2013)',
    'Certificato SSL attivo (connessione HTTPS)',
    'Nessun tracker di marketing (no Google Analytics, no Facebook Pixel)',
    'Nessun embed problematico (no YouTube, no widget social)',
    'Checkbox privacy presente nel form contatti',
]
for item in positives:
    add_bullet(doc, item)

# ============================================
# 2. ASPETTI DA CORREGGERE
# ============================================
doc.add_heading('2. Aspetti da correggere', level=2)

# 2.1 Cookie Banner
doc.add_heading('2.1 Cookie banner — Priorità alta', level=3)

p = doc.add_paragraph()
add_bold_text(p, 'Situazione: ')
p.add_run('Il sito non presenta un banner per la raccolta del consenso sui cookie.')

p = doc.add_paragraph()
add_bold_text(p, 'Perché è necessario: ')
p.add_run(
    'Il sito utilizza Google reCAPTCHA, che installa cookie classificati come non strettamente tecnici. '
    'Le Linee Guida del Garante Privacy del 10 giugno 2021 (provvedimento n. 231) richiedono il consenso '
    'preventivo dell\'utente prima dell\'installazione di cookie non tecnici. Questo vale per qualsiasi sito, '
    'inclusi i siti vetrina.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Cosa deve fare il banner:')

banner_items = [
    'Mostrare pulsante "Accetta" e "Rifiuta" con la stessa evidenza visiva',
    'Offrire la gestione delle preferenze per categoria',
    'Bloccare i cookie non tecnici (come quelli di reCAPTCHA) fino al consenso',
    'Registrare la scelta dell\'utente (sia accettazione che rifiuto) con un timestamp',
    'Non riproporre il banner prima di 6 mesi dalla scelta — il Garante stabilisce che il periodo minimo '
    'entro cui non ripresentare il banner è di 6 mesi, per evitare il fenomeno del "cookie fatigue". '
    'Trascorsi i 6 mesi, il consenso può essere nuovamente richiesto.',
]
for item in banner_items:
    add_bullet(doc, item)

p = doc.add_paragraph()
add_bold_text(p, 'Come risolvere: ')
p.add_run(
    'Esistono plugin WordPress dedicati (ad esempio Complianz, Cookiebot, o soluzioni equivalenti) '
    'che gestiscono il banner, il blocco preventivo dei cookie e la registrazione del consenso in modo automatico.'
)

p = doc.add_paragraph()
p.add_run('Rif. normativo: Art. 122 D.Lgs. 196/2003; Linee Guida Garante 10/06/2021 (provv. 231); Art. 7 GDPR.').italic = True

# 2.2 Informativa Privacy
doc.add_heading('2.2 Informativa privacy — Priorità alta', level=3)
doc.add_paragraph(
    'L\'informativa attuale (www.amcovr.it/informativa-privacy/) copre diversi aspetti richiesti dal GDPR '
    'ma presenta alcune lacune rispetto all\'art. 13.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Elementi da integrare:')

add_table(doc,
    ['Elemento', 'Stato attuale', 'Cosa serve'],
    [
        ['Titolare del trattamento', 'Parziale', 'Aggiungere indirizzo completo nel testo'],
        ['Finalità del trattamento', 'Presenti', 'OK'],
        ['Base giuridica', 'Assente', 'Indicare per ogni finalità la base giuridica (art. 6 GDPR). Es: form contatto = esecuzione misure precontrattuali; cookie tecnici = legittimo interesse; cookie non tecnici = consenso'],
        ['Destinatari dei dati', 'Generico', 'Indicare le categorie: hosting provider, Google (per reCAPTCHA)'],
        ['Trasferimento dati extra-UE', 'Assente', 'Da dichiarare: reCAPTCHA trasferisce dati a server Google USA'],
        ['Periodo di conservazione', 'Vago', '"Tempo strettamente necessario" non è sufficiente; servono indicazioni concrete per ogni categoria di dati'],
        ['Diritti dell\'interessato', 'Presenti', 'OK, correttamente elencati'],
        ['Data di aggiornamento', 'Assente', 'Va indicata'],
    ],
    col_widths=[4, 3, 9.5]
)

p = doc.add_paragraph()
add_bold_text(p, 'Nota: ')
p.add_run(
    'Per uno studio di amministrazione condominiale non sussiste l\'obbligo di nominare un DPO '
    '(Responsabile della Protezione dei Dati). L\'art. 37 GDPR prevede l\'obbligo solo per enti pubblici '
    'o attività che comportano monitoraggio sistematico su larga scala o trattamento massivo di dati particolari, '
    'casi che non si applicano a uno studio di amministrazione.'
)

p = doc.add_paragraph()
p.add_run('Rif. normativo: Art. 13 GDPR; Art. 37 GDPR.').italic = True

# 2.3 Cookie Policy
doc.add_heading('2.3 Cookie policy — Priorità media', level=3)
doc.add_paragraph(
    'La cookie policy attuale (www.amcovr.it/utilizzo-cookie/) risale al 30 settembre 2018, '
    'quindi è precedente alle Linee Guida del Garante del 2021.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Cosa manca:')
missing_items = [
    'Elenco specifico dei cookie installati con nome, finalità, durata e indicazione prima/terza parte',
    'reCAPTCHA non menzionato tra i servizi di terze parti',
    'Collegamento con il meccanismo di consenso (banner)',
]
for item in missing_items:
    add_bullet(doc, item)

p = doc.add_paragraph()
add_bold_text(p, 'Soluzione: ')
p.add_run('La cookie policy può essere generata automaticamente dallo stesso strumento che gestisce il cookie banner (vedi punto 2.1).')

p = doc.add_paragraph()
p.add_run('Rif. normativo: Linee Guida Garante 10/06/2021 (provv. 231), sezione "Informativa cookie".').italic = True

# 2.4 Google reCAPTCHA
doc.add_heading('2.4 Google reCAPTCHA — Priorità media', level=3)
doc.add_paragraph(
    'I form del sito utilizzano Google reCAPTCHA v3 per la protezione anti-spam. Questo servizio raccoglie '
    'dati dell\'utente (indirizzo IP, cookie, dati comportamentali) e li trasmette a server Google.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Situazione normativa:')
add_bullet(doc, 'Dal 2 aprile 2026 Google modifica il proprio ruolo per reCAPTCHA: da titolare autonomo diventa responsabile del trattamento. '
    'Questo significa che il gestore del sito diventa il titolare dei dati raccolti da reCAPTCHA e dovrà accettare un Data Processing Agreement con Google.')
add_bullet(doc, 'Il trasferimento dati verso server USA va dichiarato nell\'informativa con indicazione delle garanzie adottate.')

p = doc.add_paragraph()
add_bold_text(p, 'Due possibilità:')

p1 = doc.add_paragraph(style='List Bullet')
run = p1.add_run('Mantenere reCAPTCHA')
run.bold = True
p1.add_run(' → dichiararlo in informativa e cookie policy, bloccarlo fino al consenso nel banner, accettare il DPA Google.')

p2 = doc.add_paragraph(style='List Bullet')
run = p2.add_run('Sostituirlo con un\'alternativa europea')
run.bold = True
p2.add_run(' (es. Friendly Captcha) → nessun cookie aggiuntivo, nessun trasferimento extra-UE, nessun consenso necessario. Per un sito vetrina è la soluzione più lineare.')

p = doc.add_paragraph()
p.add_run('Rif. normativo: Art. 44-49 GDPR (trasferimenti verso paesi terzi); Provvedimento Garante 9/06/2022 (Google Analytics e trasferimenti USA).').italic = True

# 2.5 Form di contatto
doc.add_heading('2.5 Form di contatto — Priorità bassa', level=3)
doc.add_paragraph(
    'La pagina contatti ha una checkbox privacy, che è una buona prassi. I form di richiesta preventivo '
    'dovrebbero avere lo stesso trattamento: un link visibile all\'informativa e possibilmente una checkbox '
    'di presa visione.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Precisazione: ')
p.add_run(
    'La checkbox di consenso per l\'informativa non è obbligatoria quando il form serve esclusivamente a '
    'rispondere alla richiesta dell\'utente (la base giuridica è l\'art. 6.1.b GDPR — misure precontrattuali). '
    'È tuttavia consigliabile mantenerla come documentazione della presa visione.'
)

p = doc.add_paragraph()
p.add_run('Rif. normativo: Art. 13.1 GDPR.').italic = True

# ============================================
# 3. INFORMATIVA SITO VS INFORMATIVA CONDOMINI
# ============================================
doc.add_heading('3. Informativa del sito web e informativa condominiale: due cose diverse', level=2)

doc.add_paragraph(
    'Questo punto merita un approfondimento perché genera spesso confusione. '
    'Il sito web dello studio e l\'attività di amministrazione condominiale trattano dati personali '
    'di soggetti diversi, per finalità diverse e in contesti diversi. '
    'Per questo motivo richiedono due informative distinte.'
)

p = doc.add_paragraph()
add_bold_text(p, 'In sintesi:')

add_table(doc,
    ['', 'Informativa sito web', 'Informativa attività di amministrazione'],
    [
        ['A chi si rivolge', 'Visitatori del sito, chi compila i form di contatto o preventivo', 'Condòmini degli stabili gestiti dallo studio'],
        ['Quali dati copre', 'Dati di navigazione (IP, browser), cookie, dati inseriti nei form', 'Dati anagrafici, catastali, fiscali e di contatto dei condòmini'],
        ['Per quali finalità', 'Funzionamento del sito, risposta alle richieste, protezione anti-spam', 'Gestione del condominio, contabilità, assemblee, comunicazioni condominiali'],
        ['Dove deve essere accessibile', 'Sul sito web (footer, link nei form)', 'Nelle comunicazioni ai condòmini: carta intestata, firma email, documentazione consegnata'],
        ['Link attuale', 'www.amcovr.it/informativa-privacy/', 'app.amministratoreprotetto.it/informativa'],
    ],
    col_widths=[4, 6, 6.5]
)

p = doc.add_paragraph()
add_bold_text(p, 'Perché l\'informativa condominiale non va sul sito web')
doc.add_paragraph(
    'L\'art. 13 del GDPR richiede che l\'informativa venga fornita all\'interessato "nel momento in cui i dati '
    'personali sono ottenuti" e che sia specifica per il trattamento in questione. Chi visita il sito web non è '
    'necessariamente un condòmino: è un visitatore che potenzialmente compila un form o semplicemente naviga. '
    'I suoi dati vengono trattati per finalità legate al sito (navigazione, risposta alla richiesta), '
    'non per finalità di gestione condominiale.'
)
doc.add_paragraph(
    'Inserire sul sito l\'informativa condominiale (che parla di dati catastali, assemblee, contabilità) '
    'sarebbe fuorviante per il visitatore, perché descrive trattamenti che non lo riguardano. '
    'Viceversa, l\'informativa del sito (che parla di cookie e form) non sarebbe adeguata per i condòmini '
    'i cui dati vengono trattati nell\'ambito della gestione condominiale.'
)

p = doc.add_paragraph()
add_bold_text(p, 'Dove usare ciascuna informativa:')

add_bullet(doc, 'Sul sito web: ', bold_prefix=None)
p_last = doc.paragraphs[-1]
p_last.clear()
run = p_last.add_run('Informativa sito web')
run.bold = True
p_last.add_run(' → nel footer del sito e come link nei form di contatto e preventivo.')

add_bullet(doc, 'Nelle comunicazioni ai condòmini: ', bold_prefix=None)
p_last = doc.paragraphs[-1]
p_last.clear()
run = p_last.add_run('Informativa condominiale')
run.bold = True
p_last.add_run(
    ' → nella firma delle email inviate ai condòmini, nella carta intestata dello studio, '
    'e nella documentazione consegnata (es. verbali di assemblea, comunicazioni periodiche). '
    'Il link ad app.amministratoreprotetto.it/informativa nella firma email è corretto e al posto giusto.'
)

p = doc.add_paragraph()
p.add_run('Rif. normativo: Art. 13 GDPR — obbligo di fornire l\'informativa al momento della raccolta dei dati.').italic = True

# ============================================
# 4. RIEPILOGO INTERVENTI
# ============================================
doc.add_heading('4. Riepilogo interventi', level=2)

add_table(doc,
    ['#', 'Intervento', 'Priorità'],
    [
        ['1', 'Installare cookie banner conforme alle Linee Guida Garante 2021', 'Alta'],
        ['2', 'Aggiornare informativa privacy (base giuridica, conservazione, trasferimento extra-UE, data)', 'Alta'],
        ['3', 'Aggiornare cookie policy con elenco cookie dettagliato', 'Media'],
        ['4', 'Gestire correttamente reCAPTCHA o sostituirlo con alternativa EU', 'Media'],
        ['5', 'Uniformare la checkbox privacy su tutti i form', 'Bassa'],
    ],
    col_widths=[1, 12, 3.5]
)

# ============================================
# 5. RIFERIMENTI NORMATIVI
# ============================================
doc.add_heading('5. Riferimenti normativi', level=2)

refs = [
    'Regolamento UE 2016/679 (GDPR) — artt. 6, 7, 13, 37, 44-49',
    'D.Lgs. 196/2003 (Codice Privacy italiano) — art. 122 (cookie e tracciamento)',
    'Linee Guida Garante Privacy 10 giugno 2021 (provv. n. 231) — Cookie e altri strumenti di tracciamento',
    'Provvedimento Garante 9 giugno 2022 — Utilizzo di Google Analytics e trasferimento dati verso USA',
]
for ref in refs:
    add_bullet(doc, ref)

# ============================================
# DISCLAIMER FINALE
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run(
    'Il presente documento è redatto a titolo informativo e collaborativo nell\'ambito del rapporto '
    'di consulenza in essere. Non costituisce parere legale vincolante. L\'adeguamento del sito web '
    'resta nella titolarità e responsabilità di AMCO Service SRL.'
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Midas Touch SRL — Servizio Amministratore Protetto')
run.bold = True
run.font.size = Pt(10)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run2 = p2.add_run('info@midastouch.it')
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x5a, 0x98, 0x1d)

# ============================================
# SALVA
# ============================================
doc.save(output_path)
print(f"Documento salvato: {output_path}")
